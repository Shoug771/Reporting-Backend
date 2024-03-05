from flask import Flask, request, jsonify, json, send_file
from io import StringIO
import csv
import os
from docx import Document
from flask_cors import CORS
from docx2pdf import convert
import tempfile
import shutil
import zipfile


# 'request' to access incoming request data
# 'StringIO' to create a file-like object from the content
# 'jsonify' to convert Python dictionaries to JSON responses
# 'Document' for creating and manipulating word documents


app = Flask(__name__)
CORS(app)

ALLOWED_EXTENSIONS = {'csv'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


@app.route("/api/generate_report", methods=['POST'])

def generate_report():

    try:
    # Get form data from the request
       print(request.form)
       print(request.files)

       report_type = request.form.get('report_type')
       client = request.form.get('client')
       date = request.form.get('date')
       csv_file = request.files.get('csv_file')
       #pdf_file = request.files.get('pdf_file') if report_type == 'monthly' else None

     # Check if required fields are missing
       missing_fields = []
       if report_type is None:
         missing_fields.append("Report Type")
       if client is None:
         missing_fields.append("Client")
       if date is None:
         missing_fields.append("Date")

       if csv_file is None:
         missing_fields.append("CSV File is missing.")

       elif not allowed_file(csv_file.filename):
         missing_fields.append(f"Invalid CSV File: {csv_file.filename}")


       if missing_fields:
         raise ValueError("Missing required fields: " + ", ".join(missing_fields))
       

       json_data = csv_to_json(csv_file)

       word_file = generate_word(report_type, client, date, json_data)
       pdf_file = word_to_pdf(word_file)
       
       temp_dir = tempfile.mkdtemp()

        # Save the Word and PDF files in the temporary directory
       temp_word_file = os.path.join(temp_dir, os.path.basename(word_file))
       temp_pdf_file = os.path.join(temp_dir, os.path.basename(pdf_file))
       os.rename(word_file, temp_word_file)
       os.rename(pdf_file, temp_pdf_file)

        # Zip the files in the temporary directory
       zip_file_path = os.path.join(temp_dir, 'report_files.zip')
       with zipfile.ZipFile(zip_file_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            zipf.write(temp_word_file, os.path.basename(temp_word_file))
            zipf.write(temp_pdf_file, os.path.basename(temp_pdf_file))

       response = send_file(zip_file_path, mimetype='application/zip', download_name='report_files.zip', as_attachment=True)

       return response

    except Exception as e:
        error_message = str(e)
        return jsonify({"error": error_message}), 400

    finally:
        try:
        # Remove the temporary directory and its contents
         if os.path.exists(temp_dir):
            shutil.rmtree(temp_dir)
        except Exception as e:
            print(f"Error deleting temporary directory: {e}")


def csv_to_json(csv_file):
    try:
     data = []
     csv_content = csv_file.read().decode('utf-8')
     csv_content_io = StringIO(csv_content)
     csv_reader = csv.DictReader(csv_content_io)
     for row in csv_reader:
        data.append(row)

     return json.dumps(data, indent=2)
    
    except Exception as e:
       raise ValueError("Error Converting csv to json: " + str(e))


def generate_word(report_type, client, date, json_data):
   try: 
    # Get the current working directory
    current_directory = os.getcwd()

    # Define the template file name
    template_filename = "template.docx"

    # Create the full path to the template file
    template_path = os.path.join(current_directory, template_filename)
    document  = Document(template_path)

    # Save the document with a filename based on client name and selected date
    word_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), f"{report_type} report {client} {date}.docx")
      
    document.add_page_break()
    document.add_heading("JSON Data", level=1)
    document.add_paragraph(json_data)

    document.save(word_file)

    return word_file
   
   except Exception as e:
      raise ValueError("Error generating word doc: " + str(e))

def word_to_pdf(word_file):
    try:
     pdf_file = word_file.replace('.docx', '.pdf')
     convert(word_file, pdf_file)

     return pdf_file
    
    except Exception as e: 
       raise ValueError("Error converting word to pdf: " + str(e))
   

app.run(host="0.0.0.0", port=5000, debug=True)

 